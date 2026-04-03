#!/usr/bin/env python3
"""
search_docx.py — Search for text in a folder of Word documents (.docx)

Usage:
    python3 search_docx.py "search text" /path/to/folder
    python3 search_docx.py "aunt jane" ~/Documents/recipes
    python3 search_docx.py "apple" ~/Documents/recipes --case-sensitive
"""

import argparse
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Missing dependency. Install it with:")
    print("    pip3 install python-docx")
    sys.exit(1)


def search_docx(filepath: Path, search_term: str, case_sensitive: bool) -> list[str]:
    """
    Search a single .docx file for the search term.
    Returns a list of matching paragraph snippets (empty list if no match).
    """
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"  [!] Could not open {filepath.name}: {e}")
        return []

    matches = []
    compare_term = search_term if case_sensitive else search_term.lower()

    for para in doc.paragraphs:
        text = para.text
        compare_text = text if case_sensitive else text.lower()
        if compare_term in compare_text:
            # Truncate long lines for display
            snippet = text.strip()
            if len(snippet) > 120:
                idx = compare_text.find(compare_term)
                start = max(0, idx - 40)
                end = min(len(snippet), idx + len(search_term) + 40)
                snippet = ("..." if start > 0 else "") + snippet[start:end] + ("..." if end < len(snippet) else "")
            matches.append(snippet)

    # Also search tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                compare_text = text if case_sensitive else text.lower()
                if compare_term in compare_text:
                    snippet = text.strip().replace("\n", " ")
                    if len(snippet) > 120:
                        snippet = snippet[:120] + "..."
                    matches.append(f"[table] {snippet}")

    return matches


def main():
    parser = argparse.ArgumentParser(
        description="Search for text across all .docx files in a folder."
    )
    parser.add_argument("search_term", help="Text to search for")
    parser.add_argument("folder", help="Path to folder containing .docx files")
    parser.add_argument(
        "--case-sensitive", action="store_true",
        help="Case-sensitive search (default is case-insensitive)"
    )
    parser.add_argument(
        "--recursive", "-r", action="store_true",
        help="Search subfolders recursively"
    )
    parser.add_argument(
        "--filenames-only", "-f", action="store_true",
        help="Only print filenames, not the matching lines"
    )
    args = parser.parse_args()

    folder = Path(args.folder).expanduser().resolve()
    if not folder.exists():
        print(f"Error: folder not found: {folder}")
        sys.exit(1)
    if not folder.is_dir():
        print(f"Error: not a directory: {folder}")
        sys.exit(1)

    pattern = "**/*.docx" if args.recursive else "*.docx"
    docx_files = sorted(folder.glob(pattern))

    if not docx_files:
        print(f"No .docx files found in: {folder}")
        sys.exit(0)

    print(f"Searching {len(docx_files)} file(s) for: \"{args.search_term}\"")
    print(f"Folder: {folder}")
    print(f"Case-sensitive: {args.case_sensitive}")
    print("-" * 60)

    found_count = 0
    for filepath in docx_files:
        matches = search_docx(filepath, args.search_term, args.case_sensitive)
        if matches:
            found_count += 1
            print(f"\n✓  {filepath.name}")
            if not args.filenames_only:
                for snippet in matches:
                    print(f"   → {snippet}")

    print("-" * 60)
    if found_count == 0:
        print(f'No matches found for "{args.search_term}"')
    else:
        print(f'Found in {found_count} file(s)')


if __name__ == "__main__":
    main()
